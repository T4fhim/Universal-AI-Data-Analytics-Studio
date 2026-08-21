# File: src/reports/word_exporter.py
"""Static Word (.docx) report exporter via python-docx.

Same document-assembly shape as :class:`~src.reports.pdf_exporter.
PdfReportExporter` (heading, dataset summary, one block per section),
adapted to python-docx's ``Document`` API instead of reportlab's
platypus flowables. Figures are flattened to PNG via
:func:`~src.reports.rasterize.figure_to_png_bytes`, same as the PDF
exporter — Word has no notion of an interactive widget either.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Inches

from src.core.exceptions import ServiceError
from src.core.logger import get_logger
from src.reports.base_exporter import BaseReportExporter
from src.reports.rasterize import figure_to_png_bytes
from src.reports.report_content import ReportContent

_logger = get_logger(__name__)

_CHART_WIDTH_INCHES = 6.0


class WordReportExporter(BaseReportExporter):
    """Renders a :class:`ReportContent` as a static Word document."""

    @classmethod
    def export(cls, report_content: ReportContent, output_path: Path, **kwargs) -> Path:
        document = Document()
        document.add_heading(report_content.title, level=0)
        document.add_paragraph(
            f"Dataset: {report_content.dataset_name}  |  "
            f"Generated: {report_content.generated_at}  |  "
            f"Expertise level: {report_content.expertise_level.value}"
        )

        document.add_heading("Dataset Summary", level=1)
        for key, value in report_content.dataset_summary.items():
            document.add_paragraph(f"{key}: {value}")

        for section in report_content.sections:
            cls._render_section(document, section)

        try:
            document.save(str(output_path))
        except OSError as exc:
            raise ServiceError(
                f"Failed to write Word report to {output_path}: {exc}"
            ) from exc
        _logger.info("Exported Word report to %s", output_path)
        return output_path

    @staticmethod
    def _render_section(document: Document, section) -> None:
        document.add_heading(section.stage_label, level=1)
        if section.tool_name:
            document.add_paragraph(f"Tool: {section.tool_name}  |  {section.timestamp}")
        for line in section.summary_lines:
            document.add_paragraph(line, style="List Bullet")
        if section.explanation_text:
            paragraph = document.add_paragraph(section.explanation_text)
            paragraph.style = document.styles["Intense Quote"]
        if section.figure is not None:
            image_bytes = figure_to_png_bytes(section.figure)
            if image_bytes is not None:
                document.add_picture(
                    BytesIO(image_bytes), width=Inches(_CHART_WIDTH_INCHES)
                )
