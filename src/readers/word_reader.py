# File: src/readers/word_reader.py
"""Extracts tables from Word documents (.docx) into a Dataset.

:class:`WordReader` faces the same fundamental problem as
:class:`~src.readers.pdf_reader.PdfReader` — a Word document may
contain zero, one, or many tables embedded among prose — but the
extraction itself is considerably more straightforward. ``python-docx``
exposes ``document.tables`` as a flat, already-parsed list of table
objects directly; there is no whitespace-guessing, no lattice/stream
ambiguity, and no risk of the library hallucinating a table out of
ordinary paragraph text the way testing found ``camelot``'s ``stream``
mode could do for PDFs (see :mod:`~src.readers.pdf_reader`'s module
docstring for that finding). A ``.docx`` file's tables are a real,
structural part of the document format itself — Word's own XML schema
distinguishes a table element from a paragraph element unambiguously
— so extraction here is closer in spirit to milestone 2a's readers
(read a clearly-structured source) than to ``PdfReader``'s inference
problem, even though the *contract* (multi-table, zero-tables-is-valid)
is the same as ``PdfReader``'s.

Reuses two conventions already established by other readers in this
project rather than inventing new ones:

* **Zero tables is a valid, non-error outcome** — the same contract
  :class:`~src.readers.pdf_reader.PdfReader` implements, for the same
  reason (a prose-only Word document is not a malformed file).
* **Row 1 is assumed to be the header row** — the same assumption
  :class:`~src.readers.excel_reader.ExcelReader` and
  :class:`~src.readers.pdf_reader.PdfReader` both document, for the
  same reason: nothing in a Word table's structure explicitly marks a
  row as a header (Word does support marking a table row as a
  "repeat as header row" for print pagination, but this is a
  formatting hint, not reliably set, and not a dependable signal to
  build extraction logic around).
"""

from __future__ import annotations

from pathlib import Path

import docx
import pandas as pd

from src.core.exceptions import ReaderError
from src.core.logger import get_logger
from src.readers.base_reader import BaseReader
from src.services.workspace_service import Dataset

_logger = get_logger(__name__)

_WORD_EXTENSIONS = {".docx"}


class WordReader(BaseReader):
    """Extracts tables from .docx documents, one table per read.

    Table names take the form ``"Table N"`` (1-indexed, in document
    order) — a Word table has no inherent name of its own any more
    than a PDF-extracted table does, so this follows the same
    positional-naming approach
    :class:`~src.readers.pdf_reader.PdfReader` uses, simplified since
    there is no page concept to include (a ``.docx`` document is a
    single continuous flow, not discrete pages the way a PDF is).

    Legacy ``.doc`` (the pre-2007 binary Word format) is deliberately
    out of scope — ``python-docx`` only reads the modern ``.docx``
    XML-based format, and unlike Excel's ``.xls``/``.xlsx`` split
    (where installing ``xlrd`` closed an equivalent gap in milestone
    2b), there is no comparably lightweight, well-maintained pure-Python
    library for legacy ``.doc`` extraction available in this project's
    environment. This is a real, acknowledged gap, not a silent one.
    """

    SUPPORTED_EXTENSIONS = _WORD_EXTENSIONS

    @classmethod
    def can_read(cls, path: Path) -> bool:
        return path.suffix.lower() in _WORD_EXTENSIONS

    @classmethod
    def list_tables(cls, path: Path) -> list[str]:
        """Return names for each table found in the .docx file at ``path``.

        Returns an empty list if the document genuinely contains no
        tables — a normal, valid outcome for a prose-only document,
        not an error.

        Raises:
            ReaderError: If the file does not exist or cannot be
                opened as a .docx document at all (corrupted file,
                wrong format despite the extension — including a
                legacy ``.doc`` file that was simply renamed to
                ``.docx``, which ``python-docx`` cannot open despite
                the extension matching).
        """
        if not path.exists():
            raise ReaderError(f"Word document does not exist: {path}")

        document = cls._open_document(path)
        return [f"Table {i + 1}" for i in range(len(document.tables))]

    @classmethod
    def read(cls, path: Path, table_name: str | None = None) -> Dataset:
        """Extract one table from the .docx file at ``path``.

        Args:
            path: The Word document to read.
            table_name: Which table to read, in the ``"Table N"``
                format :meth:`list_tables` returns. If the document
                has exactly one table, ``None`` reads it directly. If
                it has more than one and ``table_name`` is ``None``,
                this raises rather than guessing.

        Raises:
            ReaderError: If the file does not exist, cannot be opened,
                contains zero tables, has more than one table but no
                ``table_name`` was given, or ``table_name`` does not
                match any table in the document.
        """
        if not path.exists():
            raise ReaderError(f"Word document does not exist: {path}")

        document = cls._open_document(path)
        table_count = len(document.tables)

        if table_count == 0:
            raise ReaderError(
                f"{path} contains no tables. Check list_tables() "
                f"before calling read() to avoid this error for "
                f"documents with no tabular content."
            )

        available_names = [f"Table {i + 1}" for i in range(table_count)]

        if table_name is None:
            if table_count == 1:
                selected_index = 0
            else:
                raise ReaderError(
                    f"{path} contains {table_count} tables "
                    f"({', '.join(available_names)}); specify which "
                    f"one to read via the table_name argument."
                )
        elif table_name not in available_names:
            raise ReaderError(
                f"{path} has no table named '{table_name}'. Available "
                f"tables: {', '.join(available_names)}."
            )
        else:
            selected_index = available_names.index(table_name)

        selected_table = document.tables[selected_index]
        dataframe, warnings = cls._table_to_dataframe(selected_table, path)

        _logger.info(
            "Read Word table '%s' from %s: %d rows, %d columns, " "%d warning(s).",
            available_names[selected_index],
            path,
            len(dataframe),
            len(dataframe.columns),
            len(warnings),
        )

        return Dataset(
            name=(
                f"{path.stem} — {available_names[selected_index]}"
                if table_count > 1
                else path.stem
            ),
            dataframe=dataframe,
            source_format="docx",
            source_path=path,
            read_warnings=warnings,
        )

    @classmethod
    def _open_document(cls, path: Path):
        """Open ``path`` as a python-docx Document, wrapping any failure as ReaderError.

        Centralizes error handling for the one operation both
        :meth:`list_tables` and :meth:`read` need to perform first —
        opening the file — so the specific exception types
        ``python-docx`` (and the underlying ``zipfile``/``lxml``
        libraries it is built on, since ``.docx`` is a zipped XML
        format) might raise for a corrupted or non-Word file are
        caught in one place rather than duplicated in both methods.
        """
        try:
            return docx.Document(path)
        except Exception as exc:
            raise ReaderError(
                f"Failed to open {path} as a Word document: {exc}"
            ) from exc

    @classmethod
    def _table_to_dataframe(cls, table, path: Path) -> tuple[pd.DataFrame, list[str]]:
        """Convert a python-docx Table object into a DataFrame, using row 1 as the header.

        Returns a ``(dataframe, warnings)`` pair rather than just a
        DataFrame. The ragged-row handling below was originally
        written on the assumption that merged cells would cause a
        row's ``cells`` count to differ from the header's — tested
        directly against a real .docx with a genuine horizontal cell
        merge, and found false: ``python-docx`` reports a merged
        cell's text duplicated across each of its original cell
        positions, keeping every row's cell count equal to the
        table's column count regardless of merges. The mismatch this
        method guards against is kept anyway, as a defensive check
        against malformed or unusual table XML this project has not
        specifically tested, rather than removed on the assumption
        that no ragged case can ever occur.
        """
        if len(table.rows) == 0:
            return pd.DataFrame(), []

        all_rows = [[cell.text for cell in row.cells] for row in table.rows]
        header_row = all_rows[0]
        data_rows = all_rows[1:]

        expected_column_count = len(header_row)
        warnings: list[str] = []

        ragged_row_count = 0
        normalized_data_rows = []
        for row in data_rows:
            if len(row) != expected_column_count:
                ragged_row_count += 1
                if len(row) < expected_column_count:
                    row = row + [""] * (expected_column_count - len(row))
                else:
                    row = row[:expected_column_count]
            normalized_data_rows.append(row)

        if ragged_row_count > 0:
            warnings.append(
                f"{ragged_row_count} row(s) had a different number of "
                f"cells than the header row (likely due to merged "
                f"cells in the original table) and were padded or "
                f"truncated to fit — review the extracted data for "
                f"misaligned values."
            )

        dataframe = pd.DataFrame(normalized_data_rows, columns=header_row)
        return dataframe, warnings
